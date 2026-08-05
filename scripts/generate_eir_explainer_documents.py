"""Generate DOCX and PDF copies of the EIR explained Markdown document."""

from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak


def clean_inline(value: str) -> str:
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"\*([^*]+)\*", r"\1", value)
    return value.replace("  ", " ").strip()


def blocks(markdown: str):
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            yield "space", ""
        elif line.startswith("### "):
            yield "h3", clean_inline(line[4:])
        elif line.startswith("## "):
            yield "h2", clean_inline(line[3:])
        elif line.startswith("# "):
            yield "h1", clean_inline(line[2:])
        elif re.match(r"^\d+\. ", line):
            yield "number", clean_inline(re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            yield "bullet", clean_inline(line[2:])
        else:
            yield "body", clean_inline(line)


def make_docx(items, output: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for kind, value in items:
        if kind == "space":
            continue
        if kind == "h1":
            doc.add_heading(value, 0)
        elif kind == "h2":
            doc.add_heading(value, 1)
        elif kind == "h3":
            doc.add_heading(value, 2)
        elif kind == "bullet":
            doc.add_paragraph(value, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(value, style="List Number")
        else:
            doc.add_paragraph(value)
    doc.save(output)


def make_pdf(items, output: Path):
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCentered", parent=styles["Title"], alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name="SmallBody", parent=styles["BodyText"], fontSize=9.5, leading=13, spaceAfter=6))
    document = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm)
    story = []
    for kind, value in items:
        safe = value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind == "space":
            story.append(Spacer(1, 3))
        elif kind == "h1":
            story.append(Paragraph(safe, styles["TitleCentered"]))
        elif kind == "h2":
            story.append(Paragraph(safe, styles["Heading2"]))
        elif kind == "h3":
            story.append(Paragraph(safe, styles["Heading3"]))
        elif kind == "bullet":
            story.append(Paragraph("• " + safe, styles["SmallBody"]))
        elif kind == "number":
            story.append(Paragraph(safe, styles["SmallBody"]))
        else:
            story.append(Paragraph(safe, styles["SmallBody"]))
    document.build(story)


def main():
    source = Path(sys.argv[1])
    output_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else source.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    parsed = list(blocks(source.read_text(encoding="utf-8")))
    stem = source.stem
    make_docx(parsed, output_dir / f"{stem}.docx")
    make_pdf(parsed, output_dir / f"{stem}.pdf")


if __name__ == "__main__":
    main()
